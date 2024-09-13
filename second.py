import streamlit as st
import re
from langchain.prompts import PromptTemplate
from docx import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import json
import io
 
Gemini = st.secrets["GOOGLE_API_KEY"]
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", google_api_key=Gemini)
 

 
doc1_path = "my_own.docx "
doc2_path = "my_own2.docx"
doc3_path= "Data_license_Agreement.docx"
doc4_path= "professional_service_agreement.docx"
doc5_path= "asset_purchase_agreement.docx"
doc6_path= "SAFE2.docx"
doc7_path= "Stock_Purchase_Agreement_Startups.docx"
 
doc1 = Document(doc1_path)
doc2 = Document(doc2_path)
doc3=Document(doc3_path)
doc4=Document(doc4_path)
doc5=Document(doc5_path)
doc6=Document(doc6_path)
doc7=Document(doc7_path)
 
placeholders1 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc1.paragraphs]))
placeholders2 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc2.paragraphs]))
placeholders3 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc3.paragraphs]))
placeholders4 = re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc4.paragraphs]))
placeholders5= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc5.paragraphs]))
placeholders6= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc6.paragraphs]))
placeholders7= re.findall(r'\{\{(.*?)\}\}', ' '.join([p.text for p in doc7.paragraphs]))
 
def process_input(user_input, placeholders1, placeholders2,placeholders3,placeholders4,placeholders5,
                  placeholders6,placeholders7):
    template = """
    You are an expert in filling details in word documents using placeholders and providing definitions for legal conditions.
    Use the provided details to fill out the placeholders .
    Determine which document the details are for based on the context provided.
    Indicate which document the details belong to by specifying "Document: Master Service Agreement" or "Document: New York Agreement" or "Document: Data License Agreement"
    or "Document: asset purchase agreement" or "Document: Safe simple agreement for future Equity" or "Document: Founders stock purchase agreement" at the beginning of your response.
    If any details are missing, indicate which ones are missing.
   
    Document 1 (New York Agreement) placeholders: {placeholders1}
    Document 2 (Master Service Agreement) placeholders: {placeholders2}
    Document 3 (Data License Agreement) placeholders: {placeholders3}
    Document 4 (Professional service Agreement) placeholders: {placeholders4}
    Document 5 (Asset purchase agreement) placeholders: {placeholders5}
    Document 6 (Safe simple agreement for future Equity) placeholders : {placeholders6}
    Document 7 (Founders stock purchase agreement) placeholders : {placeholders7}
   
   
    For Master Service Agreement, please give only date, don't give the year and other matter.
   
    For each placeholder, provide the information in this format:
    PLACEHOLDER: information
 
    If a placeholder is missing information, use this format:
    PLACEHOLDER: MISSING
 
    Based on the user input, fill in the placeholders and provide definitions for relevant legal terms.
    Structure your response exactly like this, replacing the examples with actual content:
    [
      "document": "Master Service Agreement" or "New York Agreement" or "Data License Agreement" or "Professional service agreement" or "asset purchase agreement" or "Safe simple agreement for future Equity" or "Founders stock purchase agreement",
      "placeholders": [
          "PLACEHOLDER1": "Value1",
          "PLACEHOLDER2": "Value2",
          "PLACEHOLDER3": "MISSING"
        ]
    ]
 
    User input: {content}
   
    Response:
    """
   
    formatted_template = template.format(placeholders1=placeholders1,
                                         placeholders2=placeholders2,
                                         placeholders3=placeholders3,
                                         placeholders4=placeholders4,
                                         placeholders5=placeholders5,
                                         placeholders6=placeholders6,
                                         placeholders7=placeholders7,
                                         content=user_input)
    prompt = PromptTemplate(template=formatted_template)
    chain = prompt | llm
    response = chain.invoke({"content": user_input})
    return response
 
 
# Load documents and extract placeholders
 
# Initialize session state
if 'state' not in st.session_state:
    st.session_state.state = {
        'user_input': "",
        'collected_details': {},
        'placeholders': {},
        'processed': False,
        'document_generated': False,
        'final_doc': None,
        'document_type': ""
    }
   
   
def add_content_to_document(doc_path, placeholders):
    doc = Document(doc_path)
    # Replace placeholders
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if value != "MISSING":
                paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", value)    
    return doc
 
# Initialize session state
if 'state' not in st.session_state:
    st.session_state.state = {
        'user_input': "",
        'collected_details': {},
        'definitions': {},
        'placeholders': {},
        'processed': False,
        'document_generated': False,
        'final_doc': None,
        'document_type': ""
    }
 
 
user_input = st.text_area("Enter your query to fill the details:", value=st.session_state.state['user_input'])
if user_input and st.button("Process Input"):
    st.session_state.state['user_input'] = user_input
    with st.spinner("Processing your input..."):
        processed_response = process_input(user_input, placeholders1, placeholders2,placeholders3,placeholders4,placeholders5,placeholders6,placeholders7)
        processed_response = processed_response.content
        fresponse = processed_response.replace('[','{').replace(']','}')
        try:
            fresponse = fresponse.split("```json")[1].split("```")[0]
            fresponse = json.loads(fresponse)
        except:
            fresponse = json.loads(fresponse)
        st.session_state.state['document_type'] = fresponse.get("document", "")
        st.session_state.state['placeholders'] = fresponse.get("placeholders", {})
        st.session_state.state['processed'] = True
    st.success(f"Input processed successfully for {st.session_state.state['document_type']}!")
 
# Display placeholders and definitions side by side
if st.session_state.state['processed']:
    st.markdown('<div class="step-header">Step 2: Review and Update Details</div>', unsafe_allow_html=True)
    st.markdown('<p class="subheader">Missing Details</p>', unsafe_allow_html=True)
    for key, value in st.session_state.state['placeholders'].items():
        if value == "MISSING":
            user_detail = st.text_input(f"{key.replace('_', ' ')}:", key=key,
                                        value=st.session_state.state['collected_details'].get(key, ""))
            st.session_state.state['collected_details'][key] = user_detail
        else:
            st.text_input(f"{key.replace('_', ' ')}:", value=value, disabled=True)
 
    if st.button("Update Missing Details"):
        for key, value in st.session_state.state['collected_details'].items():
            if value:
                st.session_state.state['placeholders'][key] = value
        st.success("Missing Details updated successfully!")
 
    # st.markdown('<div class="step-header">Step 3: Generate and Download Document</div>', unsafe_allow_html=True)
    if st.button("Generate Final Document"):
        with st.spinner("Generating document..."):
            doc_paths = {
            "master service agreement": doc2_path,
            "new york agreement": doc1_path,
            "data license agreement": doc3_path,
            "professional service agreement": doc4_path,
            "asset purchase agreement": doc5_path,
            "safe simple agreement for future equity" : doc6_path,
            "founders stock purchase agreement" : doc7_path
        }
           
            document_type = st.session_state.state['document_type'].lower()
           
            doc_path = doc_paths.get(document_type.lower())
            if doc_path:
                st.session_state.state['final_doc'] = add_content_to_document(
                    doc_path,
                    st.session_state.state['placeholders']
                )
                st.session_state.state['document_generated'] = True
               
            else:
                st.error(f"Unknown document type: {st.session_state.state['document_type']}")
        st.success(f"Final document generated for {st.session_state.state['document_type']} with all missing details.")
    if st.session_state.state['document_generated']:
        bio = io.BytesIO()
        st.session_state.state['final_doc'].save(bio)
        st.download_button(
            label="Download Final Document",
            data=bio.getvalue(),
            file_name=f"{st.session_state.state['document_type'].replace(' ', '_').lower()}_final.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
 
    # Display the final content of placeholders and definitions
    if st.checkbox("Show final content"):
        st.json(st.session_state.state['placeholders'])
 
else:
    st.info("Please enter your query and click 'Process Input' to start.")
# st.markdown('</div>', unsafe_allow_html=True)