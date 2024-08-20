import os
import tqdm
import tempfile
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader, UnstructuredCSVLoader, TextLoader
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores.faiss import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from docx import Document
from langchain_core.documents import Document as LangchainDocument

Gemini = st.secrets["GOOGLE_API_KEY"]
llm = ChatGoogleGenerativeAI(model="gemini-1.5-pro-latest", google_api_key=Gemini)


def load_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return [LangchainDocument(page_content=text, metadata={"source": file_path})]

@st.cache_resource(show_spinner=False)
def load_file(uploaded_file, file_name):
    temp_dir = tempfile.TemporaryDirectory()
    temp_file_path = os.path.join(temp_dir.name, uploaded_file.name)
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.read())
    
    documents = []
    try:
        if uploaded_file.name.lower().endswith(".pdf"):
            loader = PyPDFLoader(temp_file_path)
            documents = loader.load_and_split()
        elif uploaded_file.name.lower().endswith(".csv"):
            loader = UnstructuredCSVLoader(temp_file_path)
            documents = loader.load_and_split()
        elif uploaded_file.name.lower().endswith(".docx"):
            documents = load_docx(temp_file_path)
        elif uploaded_file.name.lower().endswith((".xlsx", ".xls")):
            loader = UnstructuredExcelLoader(temp_file_path)
            documents = loader.load_and_split()
            
        elif uploaded_file.name.lower().endswith((".txt")):
            loader = TextLoader(temp_file_path)
            documents = loader.load_and_split()
        else:
            raise ValueError("Unsupported file type")
    except Exception as error:
        print(error)
        raise

    embeddings = HuggingFaceEmbeddings()
    vector = FAISS.from_documents(documents, embedding=embeddings)
    vector.save_local(f"{file_name}_INDEX")
    print("FAISS DB SAVED")
    return documents

@st.cache_resource(show_spinner=False)
def get_summarized_response(embedding_path, file_upload_name):
    cache_key = f"embeddings_{embedding_path}_{file_upload_name}"
    chain = None
    if cache_key not in st.session_state:
        embeddings = HuggingFaceEmbeddings()
        vector = FAISS.load_local(embedding_path, embeddings, allow_dangerous_deserialization=True)
        db = vector.as_retriever()
   
        template="""  You are an expert legal assistant specializing in case analysis and legal research. Your task is to assist lawyers by providing accurate and relevant information about their cases based on the available legal documents and case files.

        Use the following context to answer the lawyer's question:

        <context>
        {context}
        </context>

        Lawyer's Question: {input}

        Please provide a comprehensive response that includes:
        1. Relevant case details (e.g., case number, parties involved, date of filing)
        2. Key legal issues identified
        3. Applicable laws or precedents
        4. Important dates or deadlines
        5. Any critical findings or rulings
        6. Suggested next steps or strategies, if appropriate

        If the information is not available in the context, clearly state which parts of the question you cannot answer due to lack of information. If you can partially answer the question, provide the available information and indicate what is missing.

        Ensure your response is clear, concise, and directly relevant to the lawyer's query. Use legal terminology appropriately, but also explain complex concepts when necessary.
        """ 


        prompt = ChatPromptTemplate.from_template(template=template)
        doc_chain = create_stuff_documents_chain(llm, prompt)
        chain = create_retrieval_chain(db, doc_chain)
        st.session_state[cache_key] = chain
    else:
        chain = st.session_state[cache_key]
    return chain

